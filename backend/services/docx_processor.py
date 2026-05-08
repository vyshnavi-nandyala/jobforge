import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from loguru import logger

def extract_text_from_docx(file_path: str) -> str:
    try:
        doc = Document(file_path)
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        lines.append(text)
        return "\n".join(lines)
    except Exception as e:
        logger.error(f"Error extracting DOCX: {e}")
        return ""

def generate_docx_from_text(text: str, output_path: str, base_path: str = None) -> str:
    try:
        if base_path and os.path.exists(base_path):
            doc = Document(base_path)
            for para in doc.paragraphs:
                para.clear()
            if doc.paragraphs:
                doc.paragraphs[0].text = ""
        else:
            doc = Document()
            _set_default_styles(doc)

        lines = text.split("\n")
        current_para = None

        for line in lines:
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
                continue

            if _is_section_header(stripped):
                p = doc.add_paragraph()
                run = p.add_run(stripped.upper())
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
            elif stripped.startswith(("•", "-", "*", "–")):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(stripped.lstrip("•-*– "))
            else:
                p = doc.add_paragraph(stripped)

        doc.save(output_path)
        logger.info(f"Generated DOCX: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"Error generating DOCX: {e}")
        _generate_simple_docx(text, output_path)
        return output_path

def _set_default_styles(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

def _is_section_header(line: str) -> bool:
    headers = [
        "experience", "education", "skills", "summary", "objective",
        "projects", "certifications", "awards", "publications", "work experience",
        "professional experience", "technical skills", "contact",
    ]
    return any(line.lower().startswith(h) or line.lower() == h for h in headers)

def _generate_simple_docx(text: str, output_path: str):
    doc = Document()
    for line in text.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())
    doc.save(output_path)
